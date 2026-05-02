"""
Fix References section - saves to temp file then replaces.
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import shutil
import os

FILEPATH = r"c:\Users\Shorno\WebstormProjects\prepify\report_files\My Report.docx"
TEMPPATH = r"c:\Users\Shorno\WebstormProjects\prepify\report_files\My Report_temp.docx"

# Copy file first
shutil.copy2(FILEPATH, TEMPPATH)
doc = Document(TEMPPATH)

# ============================================================
# STEP 1: Fix styles of existing references
# ============================================================
print("=== FIXING REFERENCE STYLES ===")
for idx in range(377, 384):
    p = doc.paragraphs[idx]
    if p.style.name == "Normal" and p.text.strip():
        p.style = doc.styles["List Paragraph"]
        print(f"  Fixed para {idx}: Normal -> List Paragraph")

# ============================================================
# STEP 2: Remove empty paragraph between [1] and [2]
# ============================================================
print("\n=== REMOVING EMPTY PARAGRAPH ===")
p374 = doc.paragraphs[374]
if not p374.text.strip() and p374.style.name == "Body Text":
    parent = p374._element.getparent()
    parent.remove(p374._element)
    print(f"  Removed empty paragraph at index 374")

# ============================================================
# STEP 3: Add new technology references
# ============================================================
print("\n=== ADDING NEW REFERENCES ===")

new_refs = [
    '[11] Better Auth, "Authentication Framework for Next.js," 2025. [Online]. Available: https://www.better-auth.com/. [Accessed: 25-Jan-2026].',
    '[12] Drizzle Team, "Drizzle ORM - TypeScript ORM for SQL Databases," 2025. [Online]. Available: https://orm.drizzle.team/. [Accessed: 25-Jan-2026].',
    '[13] Cloudinary, "Image and Video Management Platform," Cloudinary.com. Available: https://cloudinary.com/. [Accessed: 25-Jan-2026].',
    '[14] Neon, "Serverless Postgres - Neon," Neon.tech. Available: https://neon.tech/. [Accessed: 25-Jan-2026].',
    '[15] Meta Platforms, "React - A JavaScript Library for Building User Interfaces," 2025. [Online]. Available: https://react.dev/. [Accessed: 25-Jan-2026].',
    '[16] Brevo, "Email Marketing and Transactional Email API," Brevo.com. Available: https://www.brevo.com/. [Accessed: 25-Jan-2026].',
    '[17] S. Maffeis, J. C. Mitchell, and A. Taly, "An Operational Semantics for JavaScript," in Proc. 6th Asian Symposium on Programming Languages and Systems (APLAS), Bangalore, India, 2008, pp. 307-325. DOI: 10.1007/978-3-540-89330-1_22.',
    '[18] The PostgreSQL Global Development Group, "PostgreSQL: The World\'s Most Advanced Open Source Relational Database," 2025. [Online]. Available: https://www.postgresql.org/. [Accessed: 25-Jan-2026].',
    '[19] Zustand Contributors, "Zustand - Bear Necessities for State Management in React," GitHub, 2025. [Online]. Available: https://github.com/pmndrs/zustand. [Accessed: 25-Jan-2026].',
    '[20] Shadcn, "shadcn/ui - Beautifully Designed Components Built with Radix UI and Tailwind CSS," 2025. [Online]. Available: https://ui.shadcn.com/. [Accessed: 25-Jan-2026].',
]

# Find [10] reference
last_ref_idx = None
for i, p in enumerate(doc.paragraphs):
    txt = p.text or ""
    if "[10]" in txt and "Next.js" in txt:
        last_ref_idx = i
        break

if last_ref_idx:
    print(f"  Found [10] at paragraph {last_ref_idx}")
    anchor = doc.paragraphs[last_ref_idx]._element
    
    for ref_text in reversed(new_refs):
        new_p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), 'ListParagraph')
        pPr.append(pStyle)
        new_p.append(pPr)
        
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = ref_text
        r.append(t)
        new_p.append(r)
        
        anchor.addnext(new_p)
    
    print(f"  Added {len(new_refs)} new references ([11] through [20])")

# ============================================================
# SAVE to temp file
# ============================================================
doc.save(TEMPPATH)
print(f"\n=== SAVED to {TEMPPATH} ===")

# ============================================================
# VERIFICATION
# ============================================================
print("\n=== VERIFICATION ===")
doc2 = Document(TEMPPATH)

ref_heading_idx = None
for i, p in enumerate(doc2.paragraphs):
    txt = p.text or ""
    if txt.strip() == "References" and "Heading" in p.style.name:
        ref_heading_idx = i
        break

if ref_heading_idx:
    print(f"\nReferences heading at para {ref_heading_idx}")
    print(f"\nAll references:")
    for i in range(ref_heading_idx + 1, min(ref_heading_idx + 30, len(doc2.paragraphs))):
        p = doc2.paragraphs[i]
        txt = p.text or ""
        if txt.strip():
            style_ok = "OK" if p.style.name == "List Paragraph" else f"WRONG ({p.style.name})"
            print(f"  [{style_ok}] {txt[:100]}...")

print(f"\n=== PLEASE CLOSE Word, then rename '{os.path.basename(TEMPPATH)}' to '{os.path.basename(FILEPATH)}' ===")
print("Or run: Move-Item -Force '{}' '{}'".format(TEMPPATH, FILEPATH))
print("\nDone!")

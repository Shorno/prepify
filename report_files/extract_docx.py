"""Extract text content from .docx files for analysis."""
import sys
from pathlib import Path
import docx

def extract(path: str, out_path: str):
    doc = docx.Document(path)
    lines = []
    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            lines.append("")
            continue
        if style.startswith("Heading"):
            lines.append(f"\n## [{style}] {text}\n")
        elif style.startswith("Title"):
            lines.append(f"\n# [TITLE] {text}\n")
        else:
            lines.append(text)
    # tables
    for i, table in enumerate(doc.tables):
        lines.append(f"\n--- TABLE {i+1} ---")
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " | ") for c in row.cells]
            lines.append(" | ".join(cells))
        lines.append("--- END TABLE ---\n")
    Path(out_path).write_text("\n".join(lines), encoding="utf-8")
    print(f"Wrote {out_path} ({len(lines)} lines)")

if __name__ == "__main__":
    base = Path(__file__).parent
    files = [
        ("Updated FYDP Tamplate for [Summer 2025].docx", "out_template.txt"),
        ("Best Report.docx", "out_best.txt"),
        ("Example-submission.docx", "out_example.txt"),
        ("My Report.docx", "out_mine.txt"),
        ("Phase-I Evaluation Report.docx", "out_phase1.txt"),
    ]
    for src, dst in files:
        try:
            extract(str(base / src), str(base / dst))
        except Exception as e:
            print(f"FAIL {src}: {e}")

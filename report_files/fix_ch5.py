"""
Fix Chapter 5 formatting issues in My Report_temp.docx:

Issues found:
1. Paras 316-318: Full paragraphs styled as Heading 4 — should be Body Text
   "Software Standards: ...", "Hardware Standards: ...", "Communication Standards: ..."
2. Paras 321-324: Full paragraphs styled as Heading 4 — should be Body Text
   "Impact on Life: ...", "Impact on Society and Environment: ...", 
   "Ethical Aspects: ...", "Sustainability Plan: ..."
3. Para 329: Empty Normal — should be Body Text
4. Para 338: "Mapping with Knowledge Profile" styled as Normal — should be Heading 4
5. Para 347: "Mapping with Complex Engineering Activities" styled as Normal — should be Heading 4
6. Para 357: Empty Normal — should be Body Text
"""
from docx import Document

FILEPATH = r"c:\Users\Shorno\WebstormProjects\prepify\report_files\My Report_temp.docx"
doc = Document(FILEPATH)

print("=== CHAPTER 5 FORMATTING FIXES ===\n")

# ============================================================
# FIX 1: Paras 316-318 — Content paragraphs wrongly styled as Heading 4
#         These should be Body Text (they're full paragraphs, not headings)
# ============================================================
print("--- Fix 1: Standards content paragraphs (316-318) ---")
for idx in [316, 317, 318]:
    p = doc.paragraphs[idx]
    old_style = p.style.name
    p.style = doc.styles["Body Text"]
    print(f"  [{idx}] {old_style} -> Body Text | {p.text[:60]}...")

# ============================================================
# FIX 2: Paras 321-324 — Impact content paragraphs wrongly styled as Heading 4
#         These should be Body Text
# ============================================================
print("\n--- Fix 2: Impact content paragraphs (321-324) ---")
for idx in [321, 322, 323, 324]:
    p = doc.paragraphs[idx]
    old_style = p.style.name
    p.style = doc.styles["Body Text"]
    print(f"  [{idx}] {old_style} -> Body Text | {p.text[:60]}...")

# ============================================================
# FIX 3: Para 329 — Empty Normal -> Body Text
# ============================================================
print("\n--- Fix 3: Empty Normal paragraph (329) ---")
p329 = doc.paragraphs[329]
if not p329.text.strip() and p329.style.name == "Normal":
    p329.style = doc.styles["Body Text"]
    print(f"  [329] Normal -> Body Text (empty)")

# ============================================================
# FIX 4: Para 338 — "Mapping with Knowledge Profile" Normal -> Heading 4
# ============================================================
print("\n--- Fix 4: 'Mapping with Knowledge Profile' heading (338) ---")
p338 = doc.paragraphs[338]
if "Mapping with Knowledge Profile" in (p338.text or "") and p338.style.name == "Normal":
    p338.style = doc.styles["Heading 4"]
    print(f"  [338] Normal -> Heading 4 | {p338.text}")

# ============================================================
# FIX 5: Para 347 — "Mapping with Complex Engineering Activities" Normal -> Heading 4
# ============================================================
print("\n--- Fix 5: 'Mapping with Complex Engineering Activities' heading (347) ---")
p347 = doc.paragraphs[347]
if "Mapping with Complex Engineering Activities" in (p347.text or "") and p347.style.name == "Normal":
    p347.style = doc.styles["Heading 4"]
    print(f"  [347] Normal -> Heading 4 | {p347.text}")

# ============================================================
# FIX 6: Para 357 — Empty Normal -> Body Text  
# ============================================================
print("\n--- Fix 6: Empty Normal paragraph (357) ---")
p357 = doc.paragraphs[357]
if not p357.text.strip() and p357.style.name == "Normal":
    p357.style = doc.styles["Body Text"]
    print(f"  [357] Normal -> Body Text (empty)")

# ============================================================
# SAVE
# ============================================================
doc.save(FILEPATH)
print("\n=== SAVED ===")

# ============================================================
# VERIFICATION
# ============================================================
print("\n=== VERIFICATION ===")
doc2 = Document(FILEPATH)

ch5_start = None
ch6_start = None
for i, p in enumerate(doc2.paragraphs):
    txt = p.text or ""
    if txt.strip() == "Chapter 5":
        ch5_start = i
    if txt.strip() == "Chapter 6" and ch5_start is not None:
        ch6_start = i
        break

if ch5_start and ch6_start:
    print(f"\nChapter 5: paragraphs {ch5_start} to {ch6_start - 1}")
    
    # Show all headings
    print("\nHeadings:")
    for i in range(ch5_start, ch6_start):
        p = doc2.paragraphs[i]
        txt = p.text or ""
        if "Heading" in p.style.name and txt.strip():
            indent = "    " if "4" in p.style.name else "  "
            print(f"{indent}[{i}] {p.style.name}: {txt[:70]}")
    
    # Check for remaining Normal-styled issues
    normal_issues = []
    for i in range(ch5_start, ch6_start):
        p = doc2.paragraphs[i]
        txt = p.text or ""
        if p.style.name == "Normal" and txt.strip() and txt.strip() != "Chapter 5":
            normal_issues.append((i, txt[:60]))
    
    if normal_issues:
        print(f"\nWARNING: {len(normal_issues)} Normal-styled paragraphs remain:")
        for idx, txt in normal_issues:
            print(f"  [{idx}] {txt}")
    else:
        print(f"\n  No Normal-styled content issues!")
    
    # Check for Heading 4 paragraphs that look like content (long text)
    heading4_issues = []
    for i in range(ch5_start, ch6_start):
        p = doc2.paragraphs[i]
        txt = p.text or ""
        if p.style.name == "Heading 4" and len(txt) > 80:
            heading4_issues.append((i, txt[:60]))
    
    if heading4_issues:
        print(f"\nWARNING: {len(heading4_issues)} Heading 4 paragraphs with long text (may be content):")
        for idx, txt in heading4_issues:
            print(f"  [{idx}] {txt}...")
    else:
        print(f"  No Heading 4 paragraphs with suspicious long text!")

print("\nDone!")
